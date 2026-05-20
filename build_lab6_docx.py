from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = r"C:\Users\hongv\Documents\Codex\2026-05-17\t-o-1-b-i-lab5\huong-dan-lab6-audio-spectrogram-dwt-extract.docx"


def set_run_font(run, name="Arial", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=8, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, name="Consolas", size=10)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F4F6")
    p_pr.append(shd)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    if level == 1:
        run = p.add_run(text)
        set_run_font(run, size=16, bold=True, color="1F2937")
        set_paragraph_spacing(p, before=6, after=12, line=1.0)
    elif level == 2:
        run = p.add_run(text)
        set_run_font(run, size=13, bold=True, color="1F2937")
        set_paragraph_spacing(p, before=10, after=6, line=1.0)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11, bold=True, color="374151")
        set_paragraph_spacing(p, before=8, after=4, line=1.0)


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run)
    set_paragraph_spacing(p)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run)
    set_paragraph_spacing(p, after=4)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    set_run_font(run)
    set_paragraph_spacing(p, after=4)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Huong dan thuc hanh Lab 6: audio-spectrogram-dwt-extract")
    set_run_font(run, size=18, bold=True, color="111827")
    set_paragraph_spacing(title, after=14, line=1.0)

    add_body(doc, "Tai lieu nay ap dung cho Lab 6: audio-spectrogram-dwt-extract.")
    add_body(doc, "Lab dung hai container sender va receiver. Sender gui stego.wav va secret.key, receiver tach lai anh bi mat tu audio stego.")

    add_heading(doc, "Tai Bai Lab", 2)
    add_code_block(doc, "imodule https://raw.githubusercontent.com/hongvanveo/audio-spectrogram-dwt-extract/main/imodule_audio-spectrogram-dwt-extract.tar")

    add_heading(doc, "Khoi Dong", 2)
    add_code_block(doc, "labtainer -r audio-spectrogram-dwt-extract")
    add_body(doc, "Khi duoc hoi e-mail/student id, sinh vien nhap ma sinh vien cua minh. He thong se tu chuan hoa ma do sang dang IN HOA va ghi nho ID gan nhat.")
    add_body(doc, "Lenh checkwork chi hien thi va cham ket qua cua dung ID dang duoc dung cho bai lab hien tai, khong tron voi cac file .lab cu cua ID khac.")

    add_heading(doc, "Muc tieu bai lab", 2)
    add_number(doc, "Kiem tra sender co stego.wav va secret.key.")
    add_number(doc, "Bat SSH tren receiver.")
    add_number(doc, "Co the nghe truc tiep stego.wav o sender.")
    add_number(doc, "Gui stego.wav va secret.key tu sender sang receiver.")
    add_number(doc, "Sua extract_task.py de dien dung ten file dau vao.")
    add_number(doc, "Chay chuong trinh tach anh de tao recovered_secret.png.")
    add_number(doc, "Mo truc tiep recovered_secret.png.")
    add_number(doc, "Chay checkwork de kiem tra ket qua.")

    add_heading(doc, "Task 1: Kiem tra sender", 3)
    add_code_block(doc, "cd ~/stego\nls -l\ncat README_sender.txt")

    add_heading(doc, "Task 2: Bat SSH tren receiver", 3)
    add_code_block(doc, "sudo service ssh start\nsystemctl status ssh")

    add_heading(doc, "Task 3: Nghe truc tiep audio stego", 3)
    add_code_block(doc, "./play_stego.sh")

    add_heading(doc, "Task 4: Gui file sang receiver", 3)
    add_code_block(doc, "scp ~/stego/stego.wav ~/stego/secret.key ubuntu@receiver:~/stego/")
    add_code_block(doc, "cd ~/stego\nls -l stego.wav secret.key")

    add_heading(doc, "Task 5: Sua extract_task.py", 3)
    add_code_block(doc, "cd ~/stego\nnano extract_task.py")
    add_body(doc, "Sua hai dong TODO thanh:")
    add_code_block(doc, 'STEGO_FILE = "stego.wav"\nKEY_FILE = "secret.key"')

    add_heading(doc, "Task 6: Tach anh bi mat", 3)
    add_code_block(doc, "python3 extract_task.py\nls -l recovered_secret.png")
    add_body(doc, "Mo truc tiep anh trong receiver:")
    add_code_block(doc, "./view_recovered.sh")

    add_heading(doc, "Checkwork", 2)
    add_code_block(doc, "checkwork")
    add_body(doc, "Ket qua dung can co:")
    add_code_block(doc, "Y - audio_received\nY - key_received\nY - dwt_signal_extracted\nY - key_permutation_used\nY - secret_image_recovered\nY - recovered_image_viewed\nY - recovered_image_valid")
    add_bullet(doc, "audio_received: receiver da nhan stego.wav.")
    add_bullet(doc, "key_received: receiver da nhan secret.key.")
    add_bullet(doc, "dwt_signal_extracted: receiver da tach tin hieu bi mat tu he so detail cua DWT.")
    add_bullet(doc, "key_permutation_used: receiver da dung key de dao hoan vi anh.")
    add_bullet(doc, "secret_image_recovered: da tao recovered_secret.png.")
    add_bullet(doc, "recovered_image_viewed: da mo truc tiep recovered_secret.png trong lab.")
    add_bullet(doc, "recovered_image_valid: recovered_secret.png dung voi anh goc.")

    add_heading(doc, "Ket thuc", 2)
    add_code_block(doc, "stoplab audio-spectrogram-dwt-extract")
    add_body(doc, "Ket qua duoc luu trong /home/student/labtainer_xfer/audio-spectrogram-dwt-extract va file bai lam co dang B22DCAT311.audio-spectrogram-dwt-extract.lab.")

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_doc()
